"""
Converte relatorio.docx → relatorio.md (inverso de md_to_docx.py).
Detecta estilos Word e reconstrói a sintaxe Markdown correspondente.
"""

import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _shd_fill(para) -> str:
    """Devolve a cor de fundo (fill) do parágrafo, ou '' se não tiver."""
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return ""
    shd = pPr.find(qn("w:shd"))
    if shd is None:
        return ""
    return shd.get(qn("w:fill"), "")


def _is_code_block(para) -> bool:
    fill = _shd_fill(para).upper()
    if fill in ("F2F2F2", "F0F0F0"):
        return True
    # também detecta pelo nome da fonte do primeiro run
    for run in para.runs:
        if run.font.name and "Courier" in run.font.name:
            return True
    return False


def _runs_to_md(para) -> str:
    """Converte os runs de um parágrafo em texto Markdown inline."""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        is_mono = run.font.name and "Courier" in run.font.name
        if is_mono:
            parts.append(f"`{text}`")
        elif run.bold and run.italic:
            parts.append(f"***{text}***")
        elif run.bold:
            parts.append(f"**{text}**")
        elif run.italic:
            parts.append(f"*{text}*")
        else:
            parts.append(text)
    return "".join(parts)


def _table_to_md(table) -> list[str]:
    """Converte uma tabela Word em linhas Markdown."""
    lines = []
    for i, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            # texto da célula (inline MD)
            cell_text = " ".join(
                _runs_to_md(p) for p in cell.paragraphs if p.text.strip()
            ).strip()
            cells.append(cell_text)
        line = "| " + " | ".join(cells) + " |"
        lines.append(line)
        if i == 0:
            # linha separadora do cabeçalho
            sep = "| " + " | ".join("---" for _ in cells) + " |"
            lines.append(sep)
    return lines


# ---------------------------------------------------------------------------
# Converter principal
# ---------------------------------------------------------------------------

def convert(docx_path: Path, md_path: Path):
    doc = Document(str(docx_path))

    heading_map = {
        "Heading 1": "#",
        "Heading 2": "##",
        "Heading 3": "###",
        "Heading 4": "####",
        "Heading 5": "#####",
    }

    output: list[str] = []
    processed_tables: set[int] = set()

    # índice dos blocos do documento (parágrafos + tabelas intercaladas)
    body = doc.element.body
    items = list(body)

    i = 0
    while i < len(items):
        elem = items[i]
        tag = elem.tag.split("}")[-1]  # 'p' ou 'tbl'

        # ── Tabela ────────────────────────────────────────────────────────
        if tag == "tbl":
            # encontrar a tabela correspondente no doc.tables
            tbl_id = id(elem)
            if tbl_id not in processed_tables:
                processed_tables.add(tbl_id)
                # localizar no doc.tables pelo elemento XML
                for t in doc.tables:
                    if id(t._tbl) == tbl_id:
                        output.append("")
                        output.extend(_table_to_md(t))
                        output.append("")
                        break
            i += 1
            continue

        # ── Parágrafo ─────────────────────────────────────────────────────
        if tag != "p":
            i += 1
            continue

        # encontrar o objecto Paragraph correspondente
        para = None
        for p in doc.paragraphs:
            if id(p._p) == id(elem):
                para = p
                break

        if para is None:
            i += 1
            continue

        style_name = para.style.name if para.style else "Normal"
        text = para.text.strip()

        # ── Linha em branco ───────────────────────────────────────────────
        if not text:
            # evita múltiplas linhas em branco consecutivas
            if output and output[-1] != "":
                output.append("")
            i += 1
            continue

        # ── Títulos ───────────────────────────────────────────────────────
        if style_name in heading_map:
            output.append("")
            output.append(f"{heading_map[style_name]} {text}")
            output.append("")
            i += 1
            continue

        # ── Bloco de código ───────────────────────────────────────────────
        if _is_code_block(para):
            # recolhe linhas consecutivas de código
            code_lines = [para.text]
            i += 1
            while i < len(items):
                next_elem = items[i]
                next_tag = next_elem.tag.split("}")[-1]
                if next_tag != "p":
                    break
                next_para = None
                for p in doc.paragraphs:
                    if id(p._p) == id(next_elem):
                        next_para = p
                        break
                if next_para is None or not _is_code_block(next_para):
                    break
                code_lines.append(next_para.text)
                i += 1
            output.append("")
            output.append("```")
            output.extend(code_lines)
            output.append("```")
            output.append("")
            continue

        # ── Lista bullet ──────────────────────────────────────────────────
        if "List Bullet" in style_name or "List Paragraph" in style_name:
            md_text = _runs_to_md(para)
            output.append(f"- {md_text}")
            i += 1
            continue

        # ── Blockquote ────────────────────────────────────────────────────
        if para.paragraph_format.left_indent and any(
            r.font.italic for r in para.runs
        ):
            output.append(f"> {text}")
            i += 1
            continue

        # ── Parágrafo normal ──────────────────────────────────────────────
        md_text = _runs_to_md(para)
        output.append(md_text)
        i += 1

    # limpar linhas em branco múltiplas no final
    while output and output[-1] == "":
        output.pop()

    md_path.write_text("\n".join(output) + "\n", encoding="utf-8")
    print(f"Guardado: {md_path}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    base = Path(__file__).parent.parent
    docx = base / "relatorio.docx"
    md   = base / "relatorio.md"
    convert(docx, md)
