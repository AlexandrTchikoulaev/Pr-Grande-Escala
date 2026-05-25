"""
Converte relatorio.md → relatorio.docx com formatação adequada para Google Docs.
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_hyperlink_style(doc):
    """Garante que o estilo Hyperlink existe (necessário para alguns temas)."""
    styles = doc.styles
    if "Hyperlink" not in [s.name for s in styles]:
        pass


def _set_font(run, size=None, bold=None, italic=None, color=None, mono=False):
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    if mono:
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def _apply_inline(para, text):
    """Aplica negrito, itálico e inline code num parágrafo."""
    # padrão: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        # texto antes do match
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
        if m.group(1).startswith("**"):
            run = para.add_run(m.group(2))
            run.bold = True
        elif m.group(1).startswith("*"):
            run = para.add_run(m.group(3))
            run.italic = True
        elif m.group(1).startswith("`"):
            run = para.add_run(m.group(4))
            _set_font(run, mono=True, color=(100, 100, 100))
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])


def _add_table(doc, lines):
    """Converte linhas de tabela Markdown num objecto Table do docx."""
    rows = []
    for line in lines:
        if re.match(r"^\s*\|[-:| ]+\|\s*$", line):
            continue  # linha separadora
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= ncols:
                break
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            _apply_inline(p, cell_text)
            p.paragraph_format.space_after = Pt(2)
            if i == 0:
                for run in p.runs:
                    run.bold = True
                # fundo cinzento no cabeçalho
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D9D9D9")
                tcPr.append(shd)


def _add_code_block(doc, lines):
    """Adiciona um bloco de código com fundo cinzento."""
    text = "\n".join(lines)
    para = doc.add_paragraph(text)
    para.style = "Normal"
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.left_indent = Cm(0.5)
    for run in para.runs:
        _set_font(run, mono=True, color=(50, 50, 50))
    # fundo cinzento claro
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    # re-aplicar font ao run (o style pode ter sobreposto)
    for run in para.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(50, 50, 50)


def convert(md_path: Path, docx_path: Path):
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Estilos base
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]

        # ── Bloco de código ────────────────────────────────────────────────
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            _add_code_block(doc, code_lines)
            i += 1
            continue

        # ── Tabela ─────────────────────────────────────────────────────────
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            doc.add_paragraph()  # espaço depois da tabela
            continue

        # ── Título H1 ──────────────────────────────────────────────────────
        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            i += 1
            continue

        # ── H2 ─────────────────────────────────────────────────────────────
        if line.startswith("## ") and not line.startswith("### "):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            i += 1
            continue

        # ── H3 ─────────────────────────────────────────────────────────────
        if line.startswith("### ") and not line.startswith("#### "):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            i += 1
            continue

        # ── H4 ─────────────────────────────────────────────────────────────
        if line.startswith("#### "):
            text = line[5:].strip()
            p = doc.add_heading(text, level=4)
            i += 1
            continue

        # ── Linha horizontal ---  ───────────────────────────────────────────
        if re.match(r"^---+\s*$", line):
            doc.add_paragraph()
            i += 1
            continue

        # ── Blockquote > ───────────────────────────────────────────────────
        if line.startswith("> "):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            run = p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            i += 1
            continue

        # ── Lista com bullet - ─────────────────────────────────────────────
        if re.match(r"^(\s*)[-*] ", line):
            indent = len(line) - len(line.lstrip())
            text = re.sub(r"^\s*[-*] ", "", line)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.3)
            p.paragraph_format.space_after = Pt(2)
            _apply_inline(p, text)
            i += 1
            continue

        # ── Linha em branco ────────────────────────────────────────────────
        if line.strip() == "":
            i += 1
            continue

        # ── Parágrafo normal ───────────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _apply_inline(p, line.strip())
        i += 1

    doc.save(str(docx_path))
    print(f"Guardado: {docx_path}")


if __name__ == "__main__":
    base = Path(__file__).parent.parent
    md   = base / "relatorio.md"
    out  = base / "relatorio.docx"
    convert(md, out)
